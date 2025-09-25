#######################
# This program generates empty run sheets for a given consultant, and saves them all (foldered by consultant\year\month) in a given location
#######################

# Edit these. 

start_date = "2025-08-18" #MAKE SURE THIS IS A MONDAY
end_date = "2025-12-30"

base_directory = r"C:\Users\asset\OneDrive\Documents\SandB" #Put where you want the runs sheets saved here. I would recommend saving elsewhere and then cut-paste into the OneDrive. 

consultants = ["Sylvia", "Baris"]


########################

# Don't edit. If you have issues with the import lines, run setup.py. 

from docx import Document
from datetime import datetime, timedelta, date
from workalendar.oceania import Australia, Queensland, NewSouthWales, Victoria, SouthAustralia, NorthernTerritory, Tasmania, WesternAustralia, AustralianCapitalTerritory
import os
import calendar

template_path = r"C:\Users\asset\OneDrive\Desktop\Automation Stuff\Raw Programs\scheduleTemplate.docx"  # Path to your Word template
document_path = r"C:\Users\asset\OneDrive\Desktop\Automation Stuff\Raw Programs"  # Path to save the generated document

cal = Australia()
nsw = NewSouthWales()
qld = Queensland()
vic = Victoria()
sa = SouthAustralia()
nt = NorthernTerritory()
tas = Tasmania()
wa = WesternAustralia()
act = AustralianCapitalTerritory()

states = [[nsw, "NSW"], [qld, "QLD"], [vic, "VIC"], [sa, "SA"], [nt, "NT"], [tas, "TAS"], [wa, "WA"], [act, "ACT"]]

def checkHoliday(date):
    statesLs = []
    holiday = False
    holidayName = ""

    for state, name in states:
        if state.is_holiday(date):
            if not holiday:
                holiday_map = dict(state.holidays(date.year))
                holidayName = holiday_map.get(date, '?')
            holiday = True
            statesLs.append(name)

    return [holiday, statesLs, holidayName]

def generate_schedule_doc(consultant, start_date, template_path, output_path):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at {template_path}")

    fileName = f'{consultant} - Wk {start_date.day:02}_{start_date.month:02}_{start_date.year}.docx'
    fileLoc = os.path.join(output_path, fileName)
    
    doc = Document(template_path)

    for paragraph in doc.sections[0].header.paragraphs:
        if "Week of" in paragraph.text:
            for run in paragraph.runs:
                if "-consultant-" in run.text:
                    run.text = run.text.replace("-consultant-", consultant)
                if "___" in run.text:
                    run.text = run.text.replace("___", f"{start_date.day} {start_date.strftime('%B')}")

    for para in doc.paragraphs:
        for run in para.runs:
            if "-date-" in run.text:  
                holidayStat = checkHoliday(start_date)
                holidaytxt = ""
                if holidayStat[0]:
                    holidaytxt = f' - {holidayStat[2]} PH {f'({", ".join(holidayStat[1])})' if len(holidayStat[1]) != 8 else ""}'

                run.text = run.text.replace("-date-", f"{start_date.day} {start_date.strftime('%B')}" + holidaytxt)                
                
                start_date += timedelta(days=1)

    
    doc.save(fileLoc)

def futureSchedules(consultants, start_date, end_date, template_path, base_directory):
    start_date = datetime.strptime(start_date, '%Y-%m-%d').date()
    end_date = datetime.strptime(end_date, '%Y-%m-%d').date()

    for consultant in consultants:
        currentDate = start_date
        while currentDate <= end_date:

            year = currentDate.year
            month_num = currentDate.month
            month_name = calendar.month_name[currentDate.month]
            folder_path = os.path.join(base_directory, consultant, str(year), f"{month_num} {month_name}")

            os.makedirs(folder_path, exist_ok=True)

            generate_schedule_doc(consultant, currentDate, template_path, folder_path)

            currentDate += timedelta(days=7)

futureSchedules(consultants, start_date, end_date, template_path, base_directory)