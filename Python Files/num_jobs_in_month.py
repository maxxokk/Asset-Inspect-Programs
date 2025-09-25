# pip install python-docx

from __future__ import annotations
from pathlib import Path
from datetime import date
from typing import List, Dict, Tuple, Optional, Iterable
import re

# ----------------- EDIT THIS -----------------
ROOT = r"C:\Users\asset\OneDrive - Asset Inspect\AI Shared Folder"  # The folder that contains 'Runs\All'
SHOW_PER_FILE = False       # True -> print per-file counts
# --------------------------------------------

# Count "#12345" and also "#123456", etc.
PATTERN = re.compile(r"#\d{5,}")

# "... Wk DD_MM_YYYY ..." anywhere in filename (case-insensitive)
WK_DATE_RE = re.compile(r"\bWk\s*(\d{1,2})_(\d{1,2})_(\d{4})\b", re.IGNORECASE)

MONTH_NAME_TO_NUM = {
    "jan": 1, "january": 1,
    "feb": 2, "february": 2,
    "mar": 3, "march": 3,
    "apr": 4, "april": 4,
    "may": 5,
    "jun": 6, "june": 6,
    "jul": 7, "july": 7,
    "aug": 8, "august": 8,
    "sep": 9, "sept": 9, "september": 9,
    "oct": 10, "october": 10,
    "nov": 11, "november": 11,
    "dec": 12, "december": 12,
}

def _normalize_month(month: int | str) -> int:
    if isinstance(month, int):
        if 1 <= month <= 12:
            return month
        raise ValueError("Month int must be 1..12")
    s = str(month).strip()
    m = re.match(r"^\s*(\d{1,2})\b", s)
    if m:
        n = int(m.group(1))
        if 1 <= n <= 12: return n
    key = re.sub(r"[^a-z]", "", s.lower())
    if key in MONTH_NAME_TO_NUM:
        return MONTH_NAME_TO_NUM[key]
    raise ValueError(f"Unrecognised month: {month!r}")

def _parse_date_from_filename(stem: str) -> Optional[date]:
    m = WK_DATE_RE.search(stem)
    if not m: return None
    dd, mm, yyyy = map(int, m.groups())
    try:
        return date(yyyy, mm, dd)
    except ValueError:
        return None

def _iter_year_dirs(person_dir: Path, year: Optional[int]) -> Iterable[Path]:
    if year is not None:
        ydir = person_dir / str(year)
        if ydir.is_dir(): yield ydir
        return
    for d in person_dir.iterdir():
        if d.is_dir() and re.fullmatch(r"\d{4}", d.name):
            yield d

def list_job_sheets_for_month(
    root_dir: str | Path,
    month: int | str,
    year: Optional[int] = None,
) -> List[Path]:
    """
    A) Filename-first (catches misfiled docs): scan each {year} subtree and include files
       whose 'Wk DD_MM_YYYY' month matches target (and year if provided).
    B) Fallback: only inside the target '{MM Month}' folders, include files that
       lack a parsable date in the name.
    """
    root = Path(root_dir)
    base = root / "Runs" / "All"
    if not base.exists(): return []

    target_month = _normalize_month(month)
    results: set[Path] = set()
    month_globs = [f"{target_month} *", f"{target_month:02d} *"]

    for person_dir in base.iterdir():
        if not person_dir.is_dir(): continue

        for ydir in _iter_year_dirs(person_dir, year):
            # PASS A: filename date = source of truth
            for docx in ydir.rglob("*.docx"):
                d = _parse_date_from_filename(docx.stem)
                if d and d.month == target_month and (year is None or d.year == year):
                    results.add(docx)

            # PASS B: fallback — only target month folders; only undated filenames
            for g in month_globs:
                for mdir in ydir.glob(g):
                    if not mdir.is_dir(): continue
                    for docx in mdir.glob("*.docx"):
                        if _parse_date_from_filename(docx.stem) is None:
                            results.add(docx)

    def sort_key(p: Path):
        d = _parse_date_from_filename(p.stem)
        return (0, d, p.name.lower()) if d else (1, p.name.lower())

    return sorted(results, key=sort_key)

# --------------- DOCX counting ---------------
from docx import Document  # python-docx

def _counts_in_text_block(text: str) -> Tuple[int, int]:
    """
    Returns (total_matches, linewise_matches).
    linewise_matches counts at most 1 per *explicit* line (splitlines).
    """
    total = len(PATTERN.findall(text))
    linewise = sum(1 for line in text.splitlines() if PATTERN.search(line))
    return total, linewise

def _counts_in_paragraphs(paragraphs) -> Tuple[int, int]:
    t = l = 0
    for p in paragraphs:
        tt, ll = _counts_in_text_block(p.text)
        t += tt; l += ll
    return t, l

def _counts_in_tables(tables) -> Tuple[int, int]:
    t = l = 0
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                tt, ll = _counts_in_text_block(cell.text)
                t += tt; l += ll
    return t, l

def _counts_in_header_footer(hf) -> Tuple[int, int]:
    t1, l1 = _counts_in_paragraphs(hf.paragraphs)
    t2, l2 = _counts_in_tables(hf.tables)
    return (t1 + t2, l1 + l2)

def count_hash5_in_docx(path: Path) -> Tuple[int, int]:
    """
    Returns (total_matches, linewise_matches) for a .docx.
    - total_matches: every '#\\d{5,}' occurrence.
    - linewise_matches: counts at most 1 per explicit line (paragraph or soft line break).
      Word's automatic wrapping does NOT create a newline, so multiple on the same
      visual line still count as one unless there's an actual line break.
    """
    try:
        doc = Document(str(path))
    except Exception:
        return (0, 0)

    t = l = 0
    t1, l1 = _counts_in_paragraphs(doc.paragraphs); t += t1; l += l1
    t2, l2 = _counts_in_tables(doc.tables);         t += t2; l += l2
    for section in doc.sections:
        if section.header:
            tt, ll = _counts_in_header_footer(section.header); t += tt; l += ll
        if section.footer:
            tt, ll = _counts_in_header_footer(section.footer); t += tt; l += ll
    return t, l

def count_month(root_dir: str | Path, month: int | str, year: Optional[int] = None) -> Tuple[Tuple[int,int], Dict[Path, Tuple[int,int]]]:
    files = list_job_sheets_for_month(root_dir, month, year)
    per_file: Dict[Path, Tuple[int,int]] = {}
    total_all = 0
    total_linewise = 0
    for f in files:
        t, l = count_hash5_in_docx(f)
        per_file[f] = (t, l)
        total_all += t
        total_linewise += l
    return (total_all, total_linewise), per_file

# ------------------- stdin -------------------
if __name__ == "__main__":
    month_in = input("Enter month (e.g. 'July', '7', or '07 July'): ").strip()
    year_in = input("Enter year (e.g. 2025) or leave blank for ALL years: ").strip()
    yr = int(year_in) if year_in else None

    (total_all, total_linewise), per_file = count_month(ROOT, month_in, yr)

    if SHOW_PER_FILE:
        for p, (t, l) in sorted(per_file.items(), key=lambda kv: (kv[1][0], kv[0].name.lower()), reverse=True):
            print(f"{p}\n  total matches: {t}\n  per-line matches: {l}\n")
        print("-" * 60)

    label = year_in if year_in else "ALL YEARS"
    print(f"TOTAL for month={month_in} year={label}:")
    print(f"  (A) total '#'+5+ digits matches: {total_all}")
    print(f"  (B) per-line matches (at most 1 per line): {total_linewise}")
